"""
Generate HMSS SWD392 Course Project Report (.docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "C:/Users/welterial/commet/output/doc/hmss_swd392_report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set individual cell borders. kwargs: top, bottom, left, right — each a dict with keys val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            cfg = kwargs[edge]
            tag.set(qn("w:val"), cfg.get("val", "single"))
            tag.set(qn("w:sz"), str(cfg.get("sz", 4)))
            tag.set(qn("w:space"), "0")
            tag.set(qn("w:color"), cfg.get("color", "000000"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_page_number(doc: Document):
    """Add page number to footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def style_run(run, size=12, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style="Normal", bold=False, italic=False,
                  size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        style_run(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    """Add heading with Times New Roman."""
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    return p


IMAGES_DIR = "C:/Users/welterial/commet/output/doc/images"


def add_image(doc, image_path, width=Cm(16)):
    """Add an image centered in the document. Falls back to placeholder if file missing."""
    import os
    if not os.path.exists(image_path):
        add_placeholder(doc, f"[IMAGE NOT FOUND: {os.path.basename(image_path)}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(image_path, width=width)


def add_placeholder(doc, text):
    """Add a visually distinct image/content placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add border via paragraph XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "4")
        tag.set(qn("w:color"), "888888")
        pBdr.append(tag)
    pPr.append(pBdr)
    # Shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p


def make_table_style(table):
    """Apply basic table style."""
    table.style = "Table Grid"


def cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    """Set cell paragraph text cleanly."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic)
    return p


# ---------------------------------------------------------------------------
# UC table builder
# ---------------------------------------------------------------------------

UC_ROWS = [
    "UC ID and Name",
    "Created By | Date Created",
    "Primary Actor | Secondary Actors",
    "Summary",
    "Preconditions",
    "Postconditions",
    "Dependency",
    "Description of Main Sequence",
    "Description of Alternative Sequences",
    "Exceptions",
    "Nonfunctional Requirements",
    "Outstanding Questions",
]


def set_table_borders(tbl, outer_sz="12", inner_sz="6"):
    """Set table borders: thicker outer, thinner inner — matching template."""
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge, sz in [("top", outer_sz), ("left", outer_sz), ("bottom", outer_sz),
                     ("right", outer_sz), ("insideH", inner_sz), ("insideV", inner_sz)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    # Remove existing borders element if any
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_uc_table(doc, uc: dict):
    """
    Render a single use-case description table matching the template format:
    4 columns. Col 0 = label, cols 1-3 merged = value.
    Rows 1-2 use all 4 cols (label|value|label|value).
    No colored headers, no bold on labels — plain style.
    """
    # Title above table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{uc['id']}  {uc['name']}")
    style_run(run, size=11, bold=True)

    tbl = doc.add_table(rows=12, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl, outer_sz="12", inner_sz="6")

    # Column widths: label ~4cm, value ~12cm (split for 2-col rows)
    col_widths = [Cm(4.5), Cm(4.0), Cm(3.5), Cm(4.0)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    def label_value_row(row_idx, label, value):
        """Row with label in col 0, value merged across cols 1-3."""
        row = tbl.rows[row_idx]
        row.cells[1].merge(row.cells[3])
        cell_text(row.cells[0], label + ":", bold=True, size=10)
        cell_text(row.cells[1], value, size=10)

    def split_row(row_idx, label1, value1, label2, value2):
        """Row with 4 separate cells: label|value|label|value."""
        row = tbl.rows[row_idx]
        cell_text(row.cells[0], label1 + ":", bold=True, size=10)
        cell_text(row.cells[1], value1, size=10)
        cell_text(row.cells[2], label2 + ":", bold=True, size=10)
        cell_text(row.cells[3], value2, size=10)

    def multiline_value_row(row_idx, label, items):
        """Row with label in col 0, multiline value merged across cols 1-3."""
        row = tbl.rows[row_idx]
        row.cells[1].merge(row.cells[3])
        cell_text(row.cells[0], label + ":", bold=True, size=10)
        cell = row.cells[1]
        cell.text = ""
        p = cell.paragraphs[0]
        for i, item in enumerate(items):
            run = p.add_run(item)
            style_run(run, size=10)
            if i < len(items) - 1:
                p.add_run("\n")

    # Row 0: UC ID and Name (label + merged value)
    label_value_row(0, "UC ID and Name", f"{uc['id']}  {uc['name']}")

    # Row 1: Created By | Date Created (4 separate cells)
    split_row(1, "Created By", "", "Date Created", "")

    # Row 2: Primary Actor | Secondary Actors
    split_row(2, "Primary Actor", uc["primary_actor"],
              "Secondary Actors", uc["secondary_actors"])

    # Row 3-6: Summary, Preconditions, Postconditions, Dependency
    label_value_row(3, "Summary", uc["summary"])
    label_value_row(4, "Preconditions", uc["preconditions"])
    label_value_row(5, "Postconditions", uc["postconditions"])
    label_value_row(6, "Dependency", uc["dependency"])

    # Row 7: Description of main sequence
    multiline_value_row(7, "Description of main sequence", uc["main_seq"])

    # Row 8: Description of alternative sequences
    multiline_value_row(8, "Description of alternative sequences", uc["alt_seq"])

    # Row 9: Exceptions
    label_value_row(9, "Exceptions", "N/A")

    # Row 10: Nonfunctional requirements
    label_value_row(10, "Nonfunctional requirements", uc["nfr"])

    # Row 11: Outstanding questions
    label_value_row(11, "Outstanding questions", uc["questions"])

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Use-case data
# ---------------------------------------------------------------------------

USE_CASES = [
    {
        "id": "UC-01", "name": "Search Hostel Room",
        "primary_actor": "Visitor", "secondary_actors": "Google Maps",
        "summary": "Visitor searches for available hostel room listings by entering search criteria. The system returns published listings that match the entered criteria and presents the relevant listing information to the visitor.",
        "dependency": "None",
        "preconditions": "Visitor has access to the system (no login required).",
        "postconditions": "Visitor receives a list of published room listings matching the entered criteria, or is informed that no matching listing is currently available.",
        "main_seq": [
            "1. Visitor accesses the room search function.",
            "2. System displays the search form with available filter options and current published listings.",
            "3. Visitor enters or selects desired search criteria and submits the search (location, price range, amenities, availability, move-in date).",
            "4. System checks the submitted criteria.",
            "5. System returns the list of published room listings that match the submitted criteria.",
            "6. System displays the matching listings with essential information, including listing title, price, location, availability, and available map information.",
            "7. Visitor reviews the returned listings.",
        ],
        "alt_seq": [
            "Step 3: Visitor submits with no criteria — 3.1: System returns all currently published room listings using the default ordering. Continues to Step 7.",
            "Step 4: No listing matches the criteria — 4.1: System informs the visitor that no matching room is currently available. 4.2: System invites the visitor to revise the search criteria. Returns to Step 3.",
            "Step 6: Google Maps is unavailable — 6.1: System displays the matching listings without map information. Continues to Step 7.",
        ],
        "nfr": "Performance: Search results must return within 3 seconds under normal load. Usability: Filter options must be understandable to first-time users.",
        "questions": "The exact default ordering of search results when no criteria are entered will be finalized later.",
    },
    {
        "id": "UC-02", "name": "View Room Details",
        "primary_actor": "Visitor", "secondary_actors": "Google Maps",
        "summary": "Visitor selects a room listing to view its complete details. The system presents the room information, including available property and map-related information, so the visitor can review the listing.",
        "dependency": "None",
        "preconditions": "The selected room listing is publicly visible in the system.",
        "postconditions": "Visitor has viewed the room details and any available map/location information for the selected listing.",
        "main_seq": [
            "1. Visitor selects a room listing from the available listings.",
            "2. System displays the full details of the selected room listing.",
            "3. System presents relevant room information, including title, description, price, amenities, capacity, availability status, move-in date, room images, property name, property address, and basic owner information.",
            "4. Visitor reviews the displayed room details.",
            "5. Visitor requests map information for the property location.",
            "6. System displays the available map/location information for the property.",
            "7. Visitor continues reviewing the room details.",
        ],
        "alt_seq": [
            "Step 2: The selected listing is no longer publicly visible — 2.1: System informs the visitor that the room details are unavailable. 2.2: System returns the visitor to the listing results or listing page. Use case ends.",
            "Step 5: Visitor does not request map information — 5.1: System continues displaying the room details without map information. Continues to Step 7.",
            "Step 6: Map information is temporarily unavailable — 6.1: System informs the visitor that map information is temporarily unavailable. 6.2: System continues displaying the room details without map information. Continues to Step 7.",
        ],
        "nfr": "Performance: Room details must be displayed within 3 seconds under normal conditions. Usability: Room details must be readable and easy to review on standard web devices.",
        "questions": "The final set of room-detail fields shown in the first release may be refined later.",
    },
    {
        "id": "UC-03", "name": "Register Account",
        "primary_actor": "Visitor", "secondary_actors": "None",
        "summary": "Visitor creates a new Tenant or Owner account by providing required registration information. The system validates the input, creates the account with the selected role, and confirms the result.",
        "dependency": "None",
        "preconditions": "The visitor is not currently signed in.",
        "postconditions": "A new account exists in the system with the selected role and initial account status.",
        "main_seq": [
            "1. Visitor accesses the account registration function.",
            "2. System displays the registration form with available account role options (Tenant or Owner).",
            "3. Visitor selects the desired account role and enters required registration information (name, email, phone, password).",
            "4. System validates that all required fields are complete, the email is not already registered, and the password meets requirements.",
            "5. Visitor reviews the entered information and confirms the registration request.",
            "6. System creates the new account with the selected role and assigns initial account status.",
            "7. System informs the visitor that the account has been created successfully.",
        ],
        "alt_seq": [
            "Step 4: Required information is incomplete or invalid — 4.1: System informs the visitor which fields must be corrected. Returns to Step 3.",
            "Step 4: Email address is already associated with an existing account — 4.1: System informs the visitor that an account with this email already exists. Returns to Step 3.",
        ],
        "nfr": "Security: Registration credentials must be protected in transit, and passwords must be stored in a non-recoverable protected form. Performance: Registration validation must complete within an acceptable response time.",
        "questions": "Whether email verification is required immediately after registration will be finalized later.",
    },
    {
        "id": "UC-04", "name": "Sign In",
        "primary_actor": "Registered User (Tenant, Owner, or System Admin)", "secondary_actors": "None",
        "summary": "A registered user authenticates with credentials to establish an authenticated session and gain access rights based on account role.",
        "dependency": "None",
        "preconditions": "The user is not currently signed in.",
        "postconditions": "The Registered User has an authenticated session with access rights corresponding to their account role.",
        "main_seq": [
            "1. Registered User accesses the sign-in function.",
            "2. System displays the sign-in form.",
            "3. Registered User enters credentials (email and password) and submits.",
            "4. System verifies the submitted credentials and the account status.",
            "5. System establishes an authenticated session and applies access rights corresponding to the user's account role.",
            "6. System informs the Registered User that sign-in has completed successfully.",
        ],
        "alt_seq": [
            "Step 4: Credentials are invalid — 4.1: System informs the user that the credentials are incorrect. Returns to Step 3.",
            "Step 4: Account is disabled or suspended — 4.1: System informs the user that the account cannot access the system. Use case ends unsuccessfully.",
        ],
        "nfr": "Security: Credentials must be protected in transit. Authenticated sessions must be secured against hijacking. Performance: Sign-in must complete within an acceptable response time.",
        "questions": "Whether multi-factor authentication is required in release 1 will be finalized later.",
    },
    {
        "id": "UC-05", "name": "Submit Rental Request",
        "primary_actor": "Tenant", "secondary_actors": "Email Provider",
        "summary": "Tenant submits a rental request for a visible and requestable room listing by providing rental details and contact preferences. The system records the request and notifies the owner.",
        "dependency": "None",
        "preconditions": "1. Tenant has a registered account and is signed in.\n2. The selected room listing is publicly visible and in Published Available status.",
        "postconditions": "A rental request with status Pending is recorded for the selected room, visible to the Owner for review.",
        "main_seq": [
            "1. Tenant accesses the rental request function from a selected room listing.",
            "2. System displays the room summary and the rental request form.",
            "3. Tenant enters the request information: move-in date, expected rental duration, number of occupants, occupation category, budget expectation, contact phone, preferred contact method, and any special notes.",
            "4. System validates that all required fields are complete and that the room is still requestable.",
            "5. Tenant reviews the entered request information and confirms the submission.",
            "6. System records the rental request with status Pending.",
            "7. System sends a notification that a new rental request has been received to the Owner.",
            "8. System informs the Tenant that the request has been submitted successfully.",
        ],
        "alt_seq": [
            "Step 4: The room is no longer requestable (locked or hidden) — 4.1: System informs the Tenant that the request cannot be submitted for this room. Use case ends unsuccessfully.",
            "Step 4: Required request fields are incomplete or invalid — 4.1: System informs the Tenant which fields must be corrected. Returns to Step 3.",
            "Step 7: Email Provider is unavailable — 7.1: System records the notification as failed but the request submission succeeds. Continues to Step 8.",
        ],
        "nfr": "Security: Request information must be stored securely and accessible only to the relevant Owner and System Admin. Performance: Request submission must complete within an acceptable response time.",
        "questions": "The exact set of optional request fields in release 1 may be adjusted later.",
    },
    {
        "id": "UC-06", "name": "Cancel Rental Request",
        "primary_actor": "Tenant", "secondary_actors": "Email Provider",
        "summary": "Tenant cancels an eligible submitted rental request. The system updates the request status and notifies the owner.",
        "dependency": "None",
        "preconditions": "1. Tenant is signed in.\n2. Tenant has at least one submitted rental request in Pending status.",
        "postconditions": "Selected request status is Cancelled by Tenant.",
        "main_seq": [
            "1. Tenant accesses the request management function.",
            "2. System displays the Tenant's submitted requests and their current statuses.",
            "3. Tenant selects a request to cancel.",
            "4. System checks that the selected request is still in Pending status and eligible for cancellation.",
            "5. Tenant confirms the cancellation.",
            "6. System updates the request status to Cancelled by Tenant.",
            "7. System sends a notification of the cancellation to the Owner.",
            "8. System informs the Tenant that the cancellation has completed successfully.",
        ],
        "alt_seq": [
            "Step 4: Request is no longer eligible for cancellation (already accepted or rejected) — 4.1: System informs the Tenant the request cannot be cancelled. Use case ends unsuccessfully.",
            "Step 7: Email Provider unavailable — 7.1: System records the notification as failed but the cancellation still succeeds. Continues to Step 8.",
        ],
        "nfr": "Performance: Cancellation result must be reflected promptly in the request status.",
        "questions": "The detailed cancellation rule for requests already in Accepted status will be finalized later.",
    },
    {
        "id": "UC-07", "name": "Track Rental Request Status",
        "primary_actor": "Tenant", "secondary_actors": "None",
        "summary": "Tenant reviews current statuses of submitted rental requests to decide next actions.",
        "dependency": "None",
        "preconditions": "1. Tenant is signed in.",
        "postconditions": "Tenant has been informed of current rental request status, or informed that no submitted requests exist.",
        "main_seq": [
            "1. Tenant accesses the rental request status function.",
            "2. System displays the request list with current status of each request (Pending, Accepted, Rejected, Cancelled, Revoked).",
            "3. Tenant selects a specific request for more detail.",
            "4. System displays the selected request's full status details and available actions.",
            "5. Tenant reviews the detailed status information.",
        ],
        "alt_seq": [
            "Step 2: Tenant has no submitted requests — 2.1: System informs the Tenant that no request history is available. Use case ends successfully.",
        ],
        "nfr": "Performance: Status information must be accurate and up to date at time of retrieval.",
        "questions": "The exact level of request history detail visible to tenants will be finalized later.",
    },
    {
        "id": "UC-08a", "name": "Create Property",
        "primary_actor": "Owner", "secondary_actors": "None",
        "summary": "Owner creates a new property record with name, address, description, and policies. The system records the property under the owner's account.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.",
        "postconditions": "A new property record exists under the owner's account.",
        "main_seq": [
            "1. Owner accesses the create property function.",
            "2. System displays the property creation form.",
            "3. Owner enters property information: name, address, map location, description, and general policies.",
            "4. System validates that required property fields are complete.",
            "5. Owner confirms the creation.",
            "6. System records the new property.",
            "7. System informs the Owner that the property has been created successfully.",
        ],
        "alt_seq": [
            "Step 4: Required property information is incomplete or invalid — 4.1: System informs the Owner which fields must be corrected. Returns to Step 3.",
        ],
        "nfr": "Performance: Property creation must be reflected promptly after confirmation.",
        "questions": "The final list of optional property attributes may be adjusted later.",
    },
    {
        "id": "UC-08b", "name": "Update Property",
        "primary_actor": "Owner", "secondary_actors": "None",
        "summary": "Owner updates information on an existing property (name, address, description, policies). The system records the changes.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.\n2. At least one property owned by this owner exists in the system.",
        "postconditions": "The selected property's information is updated in the system.",
        "main_seq": [
            "1. Owner accesses the property management function.",
            "2. System displays the list of existing properties owned by this owner.",
            "3. Owner selects a property and requests to update it.",
            "4. System displays the current property information in an editable form.",
            "5. Owner edits property information: name, address, map location, description, and general policies.",
            "6. System validates that required property fields are complete.",
            "7. Owner confirms the update.",
            "8. System records the updated property information.",
            "9. System informs the Owner that the property has been updated successfully.",
        ],
        "alt_seq": [
            "Step 6: Required property information is incomplete or invalid — 6.1: System informs the Owner which fields must be corrected. Returns to Step 5.",
        ],
        "nfr": "Performance: Updated property information must be reflected promptly after confirmation.",
        "questions": "The final list of optional property attributes may be adjusted later.",
    },
    {
        "id": "UC-09", "name": "Create Room Listing",
        "primary_actor": "Owner", "secondary_actors": "Cloud Storage",
        "summary": "Owner creates a new room listing under an existing property, entering room details and uploading images. The system saves the listing as a draft.",
        "dependency": "None",
        "preconditions": "Owner is signed in and has at least one existing property.",
        "postconditions": "A new room listing with status Draft exists under the selected property.",
        "main_seq": [
            "1. Owner accesses the room listing creation function under a selected property.",
            "2. System displays the room listing form with required and optional fields.",
            "3. Owner enters room information: title, description, price, capacity, amenities, available-from date, furnished status, private WC status, and uploads room images.",
            "4. System validates that required fields are present and stores the uploaded images.",
            "5. Owner reviews the entered listing information and confirms saving.",
            "6. System records the new room listing with status Draft.",
            "7. System informs the Owner that the room listing has been saved as draft successfully.",
        ],
        "alt_seq": [
            "Step 4: Required fields are incomplete — 4.1: System informs the Owner which fields must be corrected. Returns to Step 3.",
            "Step 4: Cloud Storage is unavailable — 4.1: System informs the Owner that images could not be uploaded at this time. Returns to Step 3.",
        ],
        "nfr": "Performance: Listing draft must be saved reliably. Image upload must complete within an acceptable time.",
        "questions": "The final set of required versus optional room fields for draft saving will be finalized later.",
    },
    {
        "id": "UC-10", "name": "Update Room Listing",
        "primary_actor": "Owner", "secondary_actors": "Cloud Storage",
        "summary": "Owner updates information on an existing room listing. System records the updated room listing information.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.\n2. Owner has at least one existing room listing.",
        "postconditions": "Updated room listing information is recorded in the system.",
        "main_seq": [
            "1. Owner accesses the room listing management function and selects an existing listing to update.",
            "2. System displays the current room listing information in editable form.",
            "3. Owner modifies the desired fields (title, description, price, capacity, amenities, available-from date) and optionally provides new images.",
            "4. System records any new images and prepares the updated listing information for review.",
            "5. Owner reviews the changes and confirms the update.",
            "6. System records the updated room listing information.",
            "7. System informs the Owner that the listing has been updated successfully.",
        ],
        "alt_seq": [
            "Step 4: Required fields are cleared or invalid — 4.1: System informs the Owner which fields must be corrected. Returns to Step 3.",
            "Step 4: Image storage unavailable — 4.1: System informs the Owner that image storage failed at this time. Returns to Step 3.",
        ],
        "nfr": "Performance: Updated listing information must be recorded reliably and reflect promptly.",
        "questions": "Whether updating a published listing triggers a re-verification step will be finalized later.",
    },
    {
        "id": "UC-11", "name": "Publish Room Listing",
        "primary_actor": "Owner", "secondary_actors": "None",
        "summary": "A verified owner publishes a prepared room listing so it becomes publicly visible and searchable.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in and has been verified by System Admin.\n2. Selected room listing exists in Draft or unpublished state.\n3. Selected room listing has at least one stored image.",
        "postconditions": "Selected room listing is in Published Available status and appears in public search results.",
        "main_seq": [
            "1. Owner accesses the publication function for a prepared listing.",
            "2. System displays the listing information and publication requirements checklist.",
            "3. Owner reviews the listing and requests publication.",
            "4. System evaluates publication eligibility, including owner verification, listing completeness, and image availability.",
            "5. Owner confirms the publication request.",
            "6. System records the room listing as Published Available.",
            "7. System informs the Owner that the listing is now publicly searchable.",
        ],
        "alt_seq": [
            "Step 4: Owner not yet verified — 4.1: System informs the Owner that unverified owners cannot publish listings. Use case ends unsuccessfully.",
            "Step 4: Required listing fields are incomplete — 4.1: System informs the Owner which fields are missing. Returns to Step 3.",
            "Step 4: No images exist for the listing — 4.1: System informs the Owner that at least one image is required before publication. Returns to Step 3.",
        ],
        "nfr": "Performance: Publication status must be reflected promptly in public search. Security: Publication rules must be enforced consistently; unverified owners must never bypass this check.",
        "questions": "The exact minimum completeness rule for release 1 publication will be finalized later.",
    },
    {
        "id": "UC-12", "name": "Change Listing Visibility",
        "primary_actor": "Owner", "secondary_actors": "None",
        "summary": "Owner hides or archives a published room listing, removing it from public search without deleting it.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.\n2. Owner has at least one published room listing.",
        "postconditions": "Selected listing is no longer visible in public search and is in Hidden or Archived status.",
        "main_seq": [
            "1. Owner accesses the listing management function and selects a published listing.",
            "2. System displays the current listing status and available visibility actions (Hide / Archive).",
            "3. Owner selects the desired visibility action.",
            "4. System evaluates whether the selected action is valid for the listing's current status.",
            "5. Owner confirms the visibility change.",
            "6. System records the listing status as Hidden or Archived.",
            "7. System informs the Owner that the listing visibility has been changed successfully.",
        ],
        "alt_seq": [
            "Step 4: Selected action is not valid for the listing's current status — 4.1: System informs the Owner that the action cannot be applied. Use case ends unsuccessfully.",
        ],
        "nfr": "Performance: Visibility change must be reflected immediately so the listing no longer appears in public search.",
        "questions": "The archive retention policy and whether archived listings can be restored will be finalized later.",
    },
    {
        "id": "UC-13", "name": "Submit Owner Verification",
        "primary_actor": "Owner", "secondary_actors": "Cloud Storage",
        "summary": "Owner submits personal identification information and supporting documents for manual review by the System Admin.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.\n2. Owner has not yet been verified, or has a previously rejected verification submission.",
        "postconditions": "Owner verification submission is recorded with status Pending Review.",
        "main_seq": [
            "1. Owner accesses the owner verification submission function.",
            "2. System displays the required verification information fields and document requirements.",
            "3. Owner enters personal information and provides supporting identification documents.",
            "4. System records the submitted documents and prepares the verification submission for review.",
            "5. Owner reviews the submission and confirms it.",
            "6. System records the owner verification submission with status Pending Review.",
            "7. System informs the Owner that the submission has been received and is pending admin review.",
        ],
        "alt_seq": [
            "Step 4: Required information or documents are missing — 4.1: System informs the Owner which items must be corrected or completed. Returns to Step 3.",
            "Step 4: Document storage is unavailable — 4.1: System informs the Owner that document storage failed at this time. Returns to Step 3.",
        ],
        "nfr": "Security: Verification documents must be stored securely and accessible only to System Admin. Performance: Submission must be recorded reliably without data loss.",
        "questions": "The final list of required verification document types will be finalized later.",
    },
    {
        "id": "UC-14", "name": "Review Rental Request",
        "primary_actor": "Owner", "secondary_actors": "Email Provider",
        "summary": "Owner reviews submitted rental requests for a room and decides to accept, reject, or keep each pending. Accepting locks the room.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.\n2. At least one rental request exists for a room managed by the Owner.",
        "postconditions": "Selected request is updated to Accepted, Rejected, or Pending. If Accepted, the room status is Locked / Not Requestable.",
        "main_seq": [
            "1. Owner accesses the rental request review function for a selected room.",
            "2. System displays all submitted requests for that room with their visible request information.",
            "3. Owner selects a request to handle.",
            "4. System displays the request details and available decision options (Accept / Reject / Keep Pending).",
            "5. Owner selects the desired decision.",
            "6. System records the decision for the selected request.",
            "7. System records the resulting request status and, if accepted, records the room as Locked / Not Requestable.",
            "8. System sends the decision notification to the tenant.",
            "9. System informs the Owner that the decision has been recorded successfully.",
        ],
        "alt_seq": [
            "Step 5: Owner keeps request pending — 5.1: System preserves Pending status. Returns to Step 2.",
            "Step 5: Owner rejects the request — 5.1: System updates status to Rejected. 5.2: Room remains requestable if no other request was accepted. Continues to Step 8.",
            "Step 8: Email Provider unavailable — 8.1: System records the notification as failed but the decision still succeeds. Continues to Step 9.",
        ],
        "nfr": "Performance: Decision must be reflected promptly in both request and room status. Security: Request information must be accessible only to the relevant Owner and System Admin.",
        "questions": "The exact comparison view for reviewing multiple requests side by side may be refined later.",
    },
    {
        "id": "UC-15", "name": "Reopen Room Listing",
        "primary_actor": "Owner", "secondary_actors": "Email Provider",
        "summary": "Owner revokes a previously accepted rental request after the offline arrangement fails and reopens the locked room.",
        "dependency": "None",
        "preconditions": "1. Owner is signed in.\n2. At least one of the Owner's rooms has a rental request in Accepted status and is currently Locked / Not Requestable.",
        "postconditions": "Selected request status is Revoked by Owner. Room status is Published Available and accepts new rental requests.",
        "main_seq": [
            "1. Owner accesses the accepted arrangement management function.",
            "2. System displays the accepted request and the corresponding locked room information.",
            "3. Owner selects the accepted arrangement to handle.",
            "4. System displays the reopen action and the business consequence (request will be Revoked, room will become requestable again).",
            "5. Owner confirms that the offline arrangement has failed and requests reopening.",
            "6. System records the accepted request as Revoked by Owner.",
            "7. System records the room as Published Available.",
            "8. System sends notification to the Tenant that the accepted request has been revoked.",
            "9. System informs the Owner that the room listing has been reopened successfully.",
        ],
        "alt_seq": [
            "Step 3: Selected request is no longer in Accepted status — 3.1: System informs the Owner that the room cannot be reopened from this request. Use case ends unsuccessfully.",
            "Step 8: Email Provider unavailable — 8.1: System records the notification as failed but the reopen action still succeeds. Continues to Step 9.",
        ],
        "nfr": "Performance: Room status must be updated promptly so new requests can be submitted without delay.",
        "questions": "The business note shown to the tenant after revocation will be finalized later.",
    },
    {
        "id": "UC-16", "name": "Review Owner Verification",
        "primary_actor": "System Admin", "secondary_actors": "Cloud Storage, Email Provider",
        "summary": "System Admin reviews a pending owner verification submission, records an approval or rejection, updates verification status, and notifies the owner.",
        "dependency": "None",
        "preconditions": "1. System Admin is signed in.\n2. At least one owner verification submission has status Pending Review.",
        "postconditions": "Selected owner verification status is Verified or Rejected.",
        "main_seq": [
            "1. System Admin accesses the owner verification review function.",
            "2. System displays pending verification submissions.",
            "3. System Admin selects a submission to review.",
            "4. System displays the submitted owner information and supporting documents.",
            "5. System Admin reviews the submission and selects Approve or Reject.",
            "6. System records the administrative decision.",
            "7. System updates the selected owner verification status to Verified or Rejected.",
            "8. System sends notification to the Owner of the verification result.",
            "9. System informs the System Admin that the review has been completed successfully.",
        ],
        "alt_seq": [
            "Step 8: Email Provider unavailable — 8.1: System records the notification as failed but the decision still succeeds. Continues to Step 9.",
        ],
        "nfr": "Security: Verification documents must remain confidential and accessible only to System Admin during review.",
        "questions": "The final rejection-reason categories presented to the admin will be finalized later.",
    },
    {
        "id": "UC-17", "name": "Manage User Account Status",
        "primary_actor": "System Admin", "secondary_actors": "Email Provider",
        "summary": "System Admin changes a user account status by applying a permitted status transition and notifies the user.",
        "dependency": "None",
        "preconditions": "1. System Admin is signed in.\n2. The target user account exists in the system.",
        "postconditions": "User account status is updated according to the permitted transition.",
        "main_seq": [
            "1. System Admin accesses the user account administration function.",
            "2. System displays user accounts and their current statuses.",
            "3. System Admin selects a user account to manage.",
            "4. System displays the current account information and the status-management actions permitted for the selected account's current status.",
            "5. System Admin selects a permitted status-management action.",
            "6. System updates the user account status according to the selected action.",
            "7. System sends notification to the user about the status change.",
            "8. System informs the System Admin that the account-management action has been applied successfully.",
        ],
        "alt_seq": [
            "Step 4: No status-management action is available for the selected account — 4.1: System informs the System Admin that no further status change is available. Use case ends unsuccessfully.",
            "Step 6: Selected action is no longer permitted — 6.1: System informs the System Admin that the action cannot be completed. Use case ends unsuccessfully.",
            "Step 7: Email Provider unavailable — 7.1: System records the notification as failed but the account status update still succeeds. Continues to Step 8.",
        ],
        "nfr": "Security: Account-management actions must be traceable and applied consistently. Suspended or disabled accounts must not retain unauthorized access.",
        "questions": "None at this stage.",
    },
    {
        "id": "UC-18", "name": "Control Listing Visibility",
        "primary_actor": "System Admin", "secondary_actors": "Email Provider",
        "summary": "System Admin disables a suspicious or policy-violating publicly visible listing, removing it from public search and notifying the Owner.",
        "dependency": "None",
        "preconditions": "1. System Admin is signed in.\n2. The target listing is currently publicly visible.",
        "postconditions": "The listing is removed from public search and is no longer visible to visitors or tenants.",
        "main_seq": [
            "1. System Admin accesses the listing administration function.",
            "2. System displays publicly visible listings with their associated information.",
            "3. System Admin selects a listing to review.",
            "4. System displays the listing details and the Disable control action.",
            "5. System Admin selects the Disable action for the listing.",
            "6. System confirms the listing-control action.",
            "7. System removes the listing from public search.",
            "8. System sends notification to the Owner that the listing has been disabled by admin action.",
            "9. System informs the System Admin that the listing-control action has been applied successfully.",
        ],
        "alt_seq": [
            "Step 8: Email Provider unavailable — 8.1: System records the notification as failed but the listing-control action still succeeds. Continues to Step 9.",
        ],
        "nfr": "Security: Admin listing-control actions must be traceable and applied immediately to protect platform users.",
        "questions": "None at this stage.",
    },
]


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Fix heading styles
    for lvl in range(1, 5):
        hstyle = doc.styles[f"Heading {lvl}"]
        hstyle.font.name = "Times New Roman"
        hstyle.font.color.rgb = RGBColor(0, 0, 0)

    add_page_number(doc)

    # -----------------------------------------------------------------------
    # COVER PAGE
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COURSE PROJECT REPORT")
    style_run(r, size=24, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Subject: SWD392")
    style_run(r, size=16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project: Hostel Management and Search System (HMSS)")
    style_run(r, size=16, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Engineering and Technology")
    style_run(r, size=13, italic=True)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # RECORD OF CHANGES
    # -----------------------------------------------------------------------
    add_heading(doc, "Record of Changes", level=1)

    tbl = doc.add_table(rows=1, cols=4)
    make_table_style(tbl)
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ["Date", "A*/M/D", "In Charge", "Change Description"]):
        cell_text(cell, txt, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cell, "D9E1F2")
    # 3 empty rows
    for _ in range(3):
        tbl.add_row()

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # I. OVERVIEW
    # -----------------------------------------------------------------------
    add_heading(doc, "I. Overview", level=1)

    add_heading(doc, "I.1 Project Information", level=2)
    tbl = doc.add_table(rows=4, cols=2)
    make_table_style(tbl)
    rows_data = [
        ("Project Name", "Hostel Management and Search System"),
        ("Project Code", "HMSS"),
        ("Group Name", "<<Group Name>>"),
        ("Software Type", "Web Application"),
    ]
    for row, (label, value) in zip(tbl.rows, rows_data):
        cell_text(row.cells[0], label, bold=True, size=11)
        cell_text(row.cells[1], value, size=11)

    doc.add_paragraph()
    add_heading(doc, "I.2 Project Team", level=2)
    tbl2 = doc.add_table(rows=1, cols=4)
    make_table_style(tbl2)
    for cell, txt in zip(tbl2.rows[0].cells, ["Full Name", "Role", "Email", "Mobile"]):
        cell_text(cell, txt, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cell, "D9E1F2")
    for _ in range(5):
        tbl2.add_row()

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # II. REQUIREMENT SPECIFICATION
    # -----------------------------------------------------------------------
    add_heading(doc, "II. Requirement Specification", level=1)

    # II.1 Problem Description
    add_heading(doc, "II.1 Problem Description", level=2)
    add_paragraph(doc, (
        "The Hostel Management and Search System (HMSS) is a web application that enables visitors to "
        "search and view hostel room listings, tenants to submit and manage rental requests, property owners "
        "to manage properties and room listings, and system administrators to oversee platform operations "
        "including owner verification and content moderation. The context diagram below illustrates the "
        "external entities and system interfaces."
    ), space_after=6)
    add_placeholder(doc, "[Context Diagram — see Section II.3]")

    # II.2 Major Features
    add_heading(doc, "II.2 Major Features", level=2)
    features = [
        ("FE-01", "Search hostel room listings using filters (location, price, amenities, availability, move-in date)"),
        ("FE-02", "View detailed room listing information with property and map data"),
        ("FE-03", "Register a Tenant or Owner account"),
        ("FE-04", "Sign in with role-based access control"),
        ("FE-05", "Submit rental request for a room listing"),
        ("FE-06", "Cancel a pending rental request"),
        ("FE-07", "Track rental request statuses"),
        ("FE-08", "Create and update property records"),
        ("FE-09", "Create room listings with image upload (draft status)"),
        ("FE-10", "Update room listing information"),
        ("FE-11", "Publish room listing (verified owners only)"),
        ("FE-12", "Change listing visibility (hide/archive)"),
        ("FE-13", "Submit owner verification documents"),
        ("FE-14", "Review and decide on rental requests (accept/reject)"),
        ("FE-15", "Reopen locked room after failed arrangement"),
        ("FE-16", "Review owner verification submissions (admin)"),
        ("FE-17", "Manage user account status (admin)"),
        ("FE-18", "Control listing visibility for policy violations (admin)"),
    ]
    for code, desc in features:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{code}: ")
        style_run(r1, bold=True, size=11)
        r2 = p.add_run(desc)
        style_run(r2, size=11)

    # II.3 Context Diagram
    add_heading(doc, "II.3 Context Diagram", level=2)
    add_placeholder(doc, "[INSERT CONTEXT DIAGRAM IMAGE HERE]")

    add_paragraph(doc, "The following describes the data exchanges between HMSS and its external entities:", size=11, space_after=4)

    exchanges = [
        ("Visitor ↔ HMSS", "Search criteria, listing selection, registration info → search results, room details, registration outcome"),
        ("Tenant ↔ HMSS", "Credentials, rental request, cancellation, status inquiry → auth result, request status/outcome"),
        ("Owner ↔ HMSS", "Credentials, property/listing data, verification, request decisions → auth result, publication result, request details, verification status"),
        ("System Admin ↔ HMSS", "Credentials, review decisions, account/listing control → verification data, user/listing data, results"),
        ("Google Maps ↔ HMSS", "Map/geolocation data ↔ location lookup request"),
        ("Cloud Storage ↔ HMSS", "Stored references/files ↔ upload/retrieval requests"),
        ("Email Provider ↔ HMSS", "Delivery status ↔ notification dispatch"),
    ]
    for entity, desc in exchanges:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"• {entity}: ")
        style_run(r1, bold=True, size=11)
        r2 = p.add_run(desc)
        style_run(r2, size=11)

    # II.4 Nonfunctional Requirements
    add_heading(doc, "II.4 Nonfunctional Requirements", level=2)

    nfr_data = [
        ("Performance", [
            "Search results (UC-01) within 3 seconds under normal load.",
            "Room detail page (UC-02) within 3 seconds.",
            "All write operations within acceptable response time.",
        ]),
        ("Security", [
            "Credentials/sessions protected against unauthorized access.",
            "Owner verification docs accessible only to System Admin.",
            "Rental request info accessible only to relevant Owner and System Admin.",
            "Publication rules enforced consistently.",
            "Account-management actions traceable and immediate.",
        ]),
        ("Availability", [
            "Public search/viewing available during normal operating hours.",
            "Protected functions consistently accessible to authenticated users.",
        ]),
        ("Scalability", [
            "Accommodate growth in listings, users, concurrent requests.",
            "Cloud Storage and Email Provider support increased volume.",
        ]),
    ]
    for category, items in nfr_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(category)
        style_run(r, size=12, bold=True)
        for item in items:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1)
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(f"• {item}")
            style_run(r2, size=11)

    # II.5 Functional Requirements
    add_heading(doc, "II.5 Functional Requirements", level=2)

    # II.5.1 Actors
    add_heading(doc, "II.5.1 Actors", level=3)
    tbl = doc.add_table(rows=1, cols=4)
    make_table_style(tbl)
    for cell, txt in zip(tbl.rows[0].cells, ["#", "Actor", "Type", "Description"]):
        cell_text(cell, txt, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cell, "D9E1F2")

    actors = [
        ("1", "Visitor", "Human (Primary)", "Initiates room search and views listings without authentication"),
        ("2", "Tenant", "Human (Primary)", "Initiates rental requests, cancellations, and request status tracking"),
        ("3", "Owner", "Human (Primary)", "Initiates property/room management, verification submission, and request review"),
        ("4", "System Admin", "Human (Primary)", "Initiates owner verification review, account management, and listing control"),
        ("5", "Google Maps", "External System (Secondary)", "Provides location data when system displays property/listing map view"),
        ("6", "Cloud Storage", "External System (Secondary)", "Stores and serves room images when listing is published or updated"),
        ("7", "Email Provider", "External System (Secondary)", "Delivers notification messages when system triggers status-change events"),
    ]
    for row_data in actors:
        row = tbl.add_row()
        for cell, txt in zip(row.cells, row_data):
            cell_text(cell, txt, size=10)

    add_paragraph(doc, "Note: Registered User generalizes Tenant, Owner, System Admin (shared Sign In behavior).",
                  size=10, italic=True, space_before=4)

    # II.5.2 Use Cases
    add_heading(doc, "II.5.2 Use Cases", level=3)
    tbl = doc.add_table(rows=1, cols=4)
    make_table_style(tbl)
    for cell, txt in zip(tbl.rows[0].cells, ["ID", "Use Case", "Actors", "Use Case Description"]):
        cell_text(cell, txt, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cell, "D9E1F2")

    uc_summary = [
        ("UC-01", "Search Hostel Room", "Visitor / Google Maps", "Visitor searches published room listings using filters"),
        ("UC-02", "View Room Details", "Visitor / Google Maps", "Visitor views full room/property/pricing/owner info"),
        ("UC-03", "Register Account", "Visitor", "Visitor creates Tenant or Owner account"),
        ("UC-04", "Sign In", "Registered User", "Authenticated session with role-based access"),
        ("UC-05", "Submit Rental Request", "Tenant / Email Provider", "Tenant submits rental request for visible room"),
        ("UC-06", "Cancel Rental Request", "Tenant / Email Provider", "Tenant cancels eligible pending request"),
        ("UC-07", "Track Rental Request Status", "Tenant", "Tenant views current request statuses"),
        ("UC-08a", "Create Property", "Owner", "Owner creates property record"),
        ("UC-08b", "Update Property", "Owner", "Owner updates existing property"),
        ("UC-09", "Create Room Listing", "Owner / Cloud Storage", "Owner creates draft room listing with images"),
        ("UC-10", "Update Room Listing", "Owner / Cloud Storage", "Owner updates room listing"),
        ("UC-11", "Publish Room Listing", "Owner", "Verified owner publishes listing"),
        ("UC-12", "Change Listing Visibility", "Owner", "Owner hides/archives published listing"),
        ("UC-13", "Submit Owner Verification", "Owner / Cloud Storage", "Owner submits verification documents"),
        ("UC-14", "Review Rental Request", "Owner / Email Provider", "Owner accepts/rejects rental requests"),
        ("UC-15", "Reopen Room Listing", "Owner / Email Provider", "Owner revokes accepted request, reopens room"),
        ("UC-16", "Review Owner Verification", "System Admin / Cloud Storage / Email Provider", "Admin approves/rejects verification"),
        ("UC-17", "Manage User Account Status", "System Admin / Email Provider", "Admin enables/suspends/disables accounts"),
        ("UC-18", "Control Listing Visibility", "System Admin / Email Provider", "Admin disables policy-violating listings"),
    ]
    for row_data in uc_summary:
        row = tbl.add_row()
        for cell, txt in zip(row.cells, row_data):
            cell_text(cell, txt, size=10)

    # II.5.2.1 Diagram(s)
    add_heading(doc, "II.5.2.1 Diagram(s)", level=4)
    add_placeholder(doc, "[INSERT USE CASE DIAGRAM(S) HERE]")

    # II.5.2.2 Use Case Descriptions
    add_heading(doc, "II.5.2.2 Use Case Descriptions", level=4)
    for uc in USE_CASES:
        add_uc_table(doc, uc)

    # II.5.3 Activity Diagram
    add_heading(doc, "II.5.3 Activity Diagram", level=3)
    add_placeholder(doc, "[INSERT ACTIVITY DIAGRAM(S) HERE]")

    # II.6 ERD
    add_heading(doc, "II.6 Entity Relationship Diagram", level=2)
    add_placeholder(doc, "[INSERT ENTITY RELATIONSHIP DIAGRAM HERE]")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # III. ANALYSIS MODELS
    # -----------------------------------------------------------------------
    add_heading(doc, "III. Analysis Models", level=1)

    add_heading(doc, "III.1 Interaction Diagrams", level=2)
    add_heading(doc, "III.1.1 Sequence Diagram", level=3)
    add_placeholder(doc, "[INSERT SEQUENCE DIAGRAM(S) HERE]")

    add_heading(doc, "III.1.2 Communication Diagram", level=3)

    # Per-UC communication diagrams with participants and image placeholder
    COMM_DIAGRAMS = [
        ("UC-01", "Search Hostel Room",
         "Visitor, VisitorUI «user interaction», SearchCoordinator «coordinator», SearchMatchingLogic «business logic», RoomListing «entity», GoogleMapsProxy «proxy», Google Maps"),
        ("UC-02", "View Room Details",
         "Visitor, VisitorUI «user interaction», RoomDetailCoordinator «coordinator», RoomListingLogic «business logic», RoomListing «entity», GoogleMapsProxy «proxy», Google Maps"),
        ("UC-03", "Register Account",
         "Visitor, RegistrationUI «user interaction», RegistrationCoordinator «coordinator», AccountRegistrationRules «business logic», UserAccount «entity»"),
        ("UC-04", "Sign In",
         "Registered User, SignInUI «user interaction», SignInCoordinator «coordinator», AuthenticationService «service», UserAccount «entity»"),
        ("UC-05", "Submit Rental Request",
         "Tenant, RentalRequestUI «user interaction», RentalRequestCoordinator «coordinator», RentalRequestRules «business logic», RoomListing «entity», RentalRequest «entity», NotificationService «service», EmailProxy «proxy», Email Provider"),
        ("UC-06", "Cancel Rental Request",
         "Tenant, RentalRequestUI «user interaction», RentalRequestCoordinator «coordinator», RentalRequestRules «business logic», RentalRequest «entity», NotificationService «service», EmailProxy «proxy», Email Provider"),
        ("UC-07", "Track Rental Request Status",
         "Tenant, RentalRequestStatusUI «user interaction», RentalRequestCoordinator «coordinator», RentalRequestStatusRules «business logic», RentalRequest «entity»"),
        ("UC-08a", "Create Property",
         "Owner, PropertyCreationUI «user interaction», PropertyCoordinator «coordinator», Property «entity»"),
        ("UC-08b", "Update Property",
         "Owner, PropertyManagementUI «user interaction», PropertyCoordinator «coordinator», Property «entity»"),
        ("UC-09", "Create Room Listing",
         "Owner, RoomListingCreationUI «user interaction», ListingManagementCoordinator «coordinator», RoomListingRules «business logic», Property «entity», RoomListing «entity», CloudStorageProxy «proxy», Cloud Storage"),
        ("UC-10", "Update Room Listing",
         "Owner, RoomListingManagementUI «user interaction», ListingManagementCoordinator «coordinator», RoomListingRules «business logic», RoomListing «entity», CloudStorageProxy «proxy», Cloud Storage"),
        ("UC-11", "Publish Room Listing",
         "Owner, RoomListingManagementUI «user interaction», ListingManagementCoordinator «coordinator», RoomListingRules «business logic», RoomListing «entity», VerificationLogic «business logic», OwnerVerification «entity»"),
        ("UC-12", "Change Listing Visibility",
         "Owner, RoomListingManagementUI «user interaction», ListingManagementCoordinator «coordinator», RoomListingRules «business logic», RoomListing «entity»"),
        ("UC-13", "Submit Owner Verification",
         "Owner, VerificationSubmissionUI «user interaction», VerificationCoordinator «coordinator», VerificationLogic «business logic», OwnerVerification «entity», CloudStorageProxy «proxy», Cloud Storage"),
        ("UC-14", "Review Rental Request",
         "Owner, RentalRequestReviewUI «user interaction», RequestReviewCoordinator «coordinator», RentalRequestLogic «business logic», RentalRequest «entity», RoomListing «entity», NotificationService «service», EmailProxy «proxy», Email Provider"),
        ("UC-15", "Reopen Room Listing",
         "Owner, ReopenArrangementUI «user interaction», RequestReviewCoordinator «coordinator», RentalRequestLogic «business logic», RentalRequest «entity», RoomListing «entity», NotificationService «service», EmailProxy «proxy», Email Provider"),
        ("UC-16", "Review Owner Verification",
         "System Admin, VerificationReviewUI «user interaction», VerificationCoordinator «coordinator», VerificationLogic «business logic», OwnerVerification «entity», CloudStorageProxy «proxy», Cloud Storage, NotificationService «service», EmailProxy «proxy», Email Provider"),
        ("UC-17", "Manage User Account Status",
         "System Admin, AdminUI «user interaction», AdminCoordinator «coordinator», UserManagementLogic «business logic», User «entity», NotificationService «service», EmailProxy «proxy», Email Provider"),
        ("UC-18", "Control Listing Visibility",
         "System Admin, AdminUI «user interaction», AdminCoordinator «coordinator», RoomListingLogic «business logic», RoomListing «entity», NotificationService «service», EmailProxy «proxy», Email Provider"),
    ]
    for uc_id, uc_name, participants in COMM_DIAGRAMS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"{uc_id}: {uc_name} — Main Sequence")
        style_run(run, size=11, bold=True)
        p2 = doc.add_paragraph()
        run2 = p2.add_run("Participants: ")
        style_run(run2, size=10, bold=True)
        run3 = p2.add_run(participants)
        style_run(run3, size=10)
        # Map UC ID to image filename
        uc_img_key = uc_id.lower().replace("-", "-")  # e.g. "UC-01" -> "uc-01"
        img_path = f"{IMAGES_DIR}/comm-{uc_img_key}.png"
        add_image(doc, img_path)

    add_heading(doc, "III.2 State Diagram", level=2)
    add_placeholder(doc, "[INSERT STATE DIAGRAM(S) HERE]")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # IV. DESIGN SPECIFICATION
    # -----------------------------------------------------------------------
    add_heading(doc, "IV. Design Specification", level=1)

    add_heading(doc, "IV.1 Integrated Communication Diagrams", level=2)

    INTEGRATED_DIAGRAMS = [
        ("DiscoveryAndSearchSubsystem", "UC-01, UC-02",
         "VisitorUI, SearchCoordinator, RoomDetailCoordinator, SearchMatchingLogic, GoogleMapsProxy",
         "Visitor, Google Maps, RoomListing (external)"),
        ("AccountAndAccessSubsystem", "UC-03, UC-04, UC-17",
         "RegistrationUI, SignInUI, AccountAdminUI, AuthenticationLogic, UserManagementLogic",
         "Visitor, Registered User, System Admin, User, Email Provider"),
        ("RentalRequestManagementSubsystem", "UC-05, UC-06, UC-07, UC-14, UC-15",
         "RentalRequestUI, RentalRequestStatusUI, RentalRequestReviewUI, ReopenArrangementUI, RentalRequestCoordinator, RequestReviewCoordinator, RentalRequestLogic",
         "Tenant, Owner, RoomListing (external), Email Provider"),
        ("PropertyManagementSubsystem", "UC-08a, UC-08b",
         "PropertyCreationUI, PropertyManagementUI, PropertyCoordinator",
         "Owner, Property"),
        ("ListingManagementAndModerationSubsystem", "UC-09, UC-10, UC-11, UC-12, UC-18",
         "RoomListingCreationUI, RoomListingManagementUI, ListingAdminUI, ListingManagementCoordinator, RoomListingLogic",
         "Owner, System Admin, Property (external), OwnerVerification (external), Cloud Storage, Email Provider"),
        ("OwnerVerificationSubsystem", "UC-13, UC-16",
         "VerificationSubmissionUI, VerificationReviewUI, VerificationCoordinator, VerificationLogic",
         "Owner, System Admin, Cloud Storage, Email Provider"),
    ]
    for subsystem, ucs, internal, external in INTEGRATED_DIAGRAMS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"{subsystem}")
        style_run(run, size=11, bold=True)
        p2 = doc.add_paragraph()
        r1 = p2.add_run("Source Use Cases: ")
        style_run(r1, size=10, bold=True)
        r2 = p2.add_run(ucs)
        style_run(r2, size=10)
        p3 = doc.add_paragraph()
        r3 = p3.add_run("Internal Objects: ")
        style_run(r3, size=10, bold=True)
        r4 = p3.add_run(internal)
        style_run(r4, size=10)
        p4 = doc.add_paragraph()
        r5 = p4.add_run("External Actors/Objects: ")
        style_run(r5, size=10, bold=True)
        r6 = p4.add_run(external)
        style_run(r6, size=10)
        # Map subsystem to integrated diagram image if available
        subsystem_img_map = {
            "DiscoveryAndSearchSubsystem": "step-2.4a-integrated.png",
        }
        img_file = subsystem_img_map.get(subsystem)
        if img_file:
            add_image(doc, f"{IMAGES_DIR}/{img_file}")
        else:
            add_placeholder(doc, f"[INSERT {subsystem} INTEGRATED COMMUNICATION DIAGRAM HERE]")

    add_heading(doc, "IV.2 System High-Level Design", level=2)
    add_placeholder(doc, "[INSERT SYSTEM HIGH-LEVEL DESIGN HERE]")

    add_heading(doc, "IV.3 Component and Package Diagram", level=2)
    add_heading(doc, "IV.3.1 Component Diagram", level=3)
    add_placeholder(doc, "[INSERT COMPONENT DIAGRAM HERE]")
    add_heading(doc, "IV.3.2 Package Diagram", level=3)
    add_placeholder(doc, "[INSERT PACKAGE DIAGRAM HERE]")

    add_heading(doc, "IV.4 Class Diagram", level=2)
    add_placeholder(doc, "[INSERT CLASS DIAGRAM HERE]")

    add_heading(doc, "IV.5 Database Design", level=2)
    add_placeholder(doc, "[INSERT DATABASE DESIGN HERE]")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # V. IMPLEMENTATION
    # -----------------------------------------------------------------------
    add_heading(doc, "V. Implementation", level=1)

    add_heading(doc, "V.1 Map Architecture to Project Structure", level=2)
    add_placeholder(doc, "[TO BE COMPLETED]")

    add_heading(doc, "V.2 Map Class Diagram and Interaction Diagram to Code", level=2)
    add_placeholder(doc, "[TO BE COMPLETED]")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # VI. APPLYING ALTERNATIVE ARCHITECTURE PATTERNS
    # -----------------------------------------------------------------------
    add_heading(doc, "VI. Applying Alternative Architecture Patterns", level=1)

    add_heading(doc, "VI.1 Applying Alternative Architecture", level=2)
    add_placeholder(doc, "[TO BE COMPLETED]")

    add_heading(doc, "VI.2 Applying Service Discovery Pattern", level=2)
    add_placeholder(doc, "[TO BE COMPLETED]")

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
