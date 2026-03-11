from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
SOURCE_DOCX = Path(r"C:/Users/welterial/Downloads/10_01_2026___a95904323cb524c0b45ee9163b770dc4..docx")
OUTPUT_DOCX = REPO_ROOT / "output" / "doc" / "hmss_requirement_modeling_filled.docx"
USE_CASE_DIR = REPO_ROOT / "260310-hmss"

PROJECT_NAME = "Hostel Management and Search System"
PROJECT_CODE = "HMSS"
GROUP_NAME = "SWD392-G1"
SOFTWARE_TYPE = "Web-based information system"
CREATED_BY = "Project Team"
DATE_CREATED = "2026-03-10"

PROBLEM_DESCRIPTION = [
    (
        "The Hostel Management and Search System is a web-based platform that connects tenants "
        "who need suitable rental rooms with owners who manage one or more hostel properties. "
        "Within the scope of this project, a property is the overall hostel or rental house and "
        "a room is the smallest rentable unit. In release 1, one room corresponds to one listing."
    ),
    (
        "Today, tenants usually search for rooms through fragmented channels such as social media "
        "posts, chat groups, personal contacts, or inconsistent listing websites. That makes room "
        "information difficult to compare, difficult to verify, and hard to track once availability "
        "changes. Owners face the opposite side of the same problem: they lack a centralized system "
        "for managing multiple properties and rooms, publishing trustworthy listings, and reviewing "
        "rental requests in a structured way."
    ),
    (
        "HMSS addresses those problems with three business focuses: searching hostel room listings, "
        "managing properties and room listings, and handling rental requests. The system supports "
        "visitors, tenants, owners, and system administrators, while keeping the final rental "
        "arrangement outside the platform. Online payment, in-system chat, electronic contracts, "
        "and full tenancy lifecycle management are outside the current scope."
    ),
]

MAJOR_FEATURES = [
    "FE-01: Search publicly available hostel room listings using filters such as location, price range, amenities, move-in date, and requestability status.",
    "FE-02: View detailed room listing information, including property context, room conditions, amenities, pricing, and owner information.",
    "FE-03: Register and sign in as a tenant or owner account to access protected functions.",
    "FE-04: Allow owners to create and manage multiple properties and multiple rooms under each property.",
    "FE-05: Allow owners to create draft room listings and maintain listing information before publication.",
    "FE-06: Allow only verified owners to publish room listings so they become publicly searchable.",
    "FE-07: Allow tenants to submit rental requests for visible and requestable room listings.",
    "FE-08: Allow tenants to cancel submitted requests and track request status over time.",
    "FE-09: Allow owners to review rental requests and decide to accept, reject, or keep them pending.",
    "FE-10: Lock a room listing immediately after request acceptance and reopen it if the offline arrangement fails.",
    "FE-11: Allow system administrators to review owner verification submissions and manage user account status.",
    "FE-12: Allow system administrators to control the public visibility of suspicious or policy-violating listings.",
]

NFRS = [
    (
        "1",
        "Performance",
        "Search response time",
        "Search results for UC-01 should return within 3 seconds under normal load.",
    ),
    (
        "2",
        "Performance",
        "Room detail load time",
        "Room detail information for UC-02 should load within 3 seconds under normal conditions.",
    ),
    (
        "3",
        "Security",
        "Credential and session protection",
        "Credentials and authenticated sessions must be protected against unauthorized access and hijacking.",
    ),
    (
        "4",
        "Security",
        "Verification document confidentiality",
        "Owner verification documents must be stored securely and accessible only to System Admin during review.",
    ),
    (
        "5",
        "Security",
        "Request access control",
        "Rental request information must be accessible only to the relevant Owner and System Admin.",
    ),
    (
        "6",
        "Reliability",
        "Business-rule enforcement",
        "Unverified owners must never be allowed to bypass the publication gate for public listings.",
    ),
    (
        "7",
        "Availability",
        "Public and protected access",
        "Public search and protected user functions should remain available during normal operating periods with minimal interruption.",
    ),
    (
        "8",
        "Scalability",
        "Growth support",
        "The system must accommodate growth in listings, users, and concurrent rental requests without fundamental business-model changes.",
    ),
    (
        "9",
        "Integration",
        "Notification support",
        "Email-based status notifications should be supported even if temporary email delivery failures must be retried later.",
    ),
]

ACTORS = [
    (
        "1",
        "Visitor",
        "An unauthenticated user who can search public room listings and view room details.",
    ),
    (
        "2",
        "Registered User",
        "A generalized actor representing Tenant, Owner, and System Admin for the shared Sign In use case.",
    ),
    (
        "3",
        "Tenant",
        "A registered user who submits rental requests, cancels requests, and tracks request status.",
    ),
    (
        "4",
        "Owner",
        "A registered user who manages properties, room listings, owner verification submission, and rental request review.",
    ),
    (
        "5",
        "System Admin",
        "An administrator who reviews owner verification, manages account status, and controls suspicious listings.",
    ),
    (
        "6",
        "Google Maps",
        "An external service that provides location data when the system displays property or listing map information.",
    ),
    (
        "7",
        "Cloud Storage",
        "An external service that stores and serves room images and owner verification documents.",
    ),
    (
        "8",
        "Email Provider",
        "An external service that delivers business notifications such as request-status and verification-status messages.",
    ),
]

ACTIVITY_ITEMS = [
    "Search Hostel Room and View Room Details",
    "Submit Rental Request and Track Rental Request Status",
    "Review Rental Request and Reopen Room Listing",
    "Publish Room Listing and Review Owner Verification",
]

ERD_ENTITIES = [
    ("1", "User", "Stores account information, role, account status, and authentication-related data."),
    ("2", "Property", "Stores the overall hostel or rental-house information managed by an owner."),
    ("3", "RoomListing", "Stores the public-facing room listing information, visibility state, and requestability state."),
    ("4", "RentalRequest", "Stores a tenant's rental request, move-in expectations, and review status."),
    ("5", "OwnerVerification", "Stores verification submissions, review outcomes, and supporting document references for owner accounts."),
    ("6", "ListingImage", "Stores image metadata associated with a room listing."),
    ("7", "Amenity", "Stores reusable amenity definitions that can be attached to room listings."),
    ("8", "RoomAmenity", "Resolves the many-to-many relationship between room listings and amenities."),
    ("9", "Notification", "Stores optional system-generated notification records for business events."),
]

ERD_RELATIONSHIPS = [
    "One Owner manages many Properties.",
    "One Property contains one or more Room Listings in release 1.",
    "One Tenant can submit many Rental Requests.",
    "One Room Listing can receive many Rental Requests over time.",
    "One Owner has at most one active Owner Verification record at a given review stage.",
]


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def set_paragraph(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    paragraph.text = text
    if style:
        paragraph.style = style
    return paragraph


def insert_paragraph_after(block, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor = block._element if isinstance(block, Paragraph) else block._tbl
    anchor.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if text or style:
        set_paragraph(paragraph, text, style)
    return paragraph


def clone_table_after(source_table: Table, target) -> Table:
    new_tbl = copy.deepcopy(source_table._tbl)
    anchor = target._element if isinstance(target, Paragraph) else target._tbl
    anchor.addnext(new_tbl)
    return Table(new_tbl, target._parent)


def ensure_rows(table: Table, count: int) -> None:
    while len(table.rows) < count:
        table.add_row()


def fill_row(row, values):
    for cell, value in zip(row.cells, values):
        cell.text = value


def clean_markdown(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        line = line.replace("**", "")
        line = re.sub(r"^>\s*", "", line)
        lines.append(line.strip())
    return "\n".join(lines).strip()


def parse_section_map(text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = None
    for line in text.splitlines():
        match = re.match(r"^##\s+(.+)$", line.strip())
        if match:
            current = match.group(1).strip()
            sections[current] = []
            continue
        if current is not None:
            sections[current].append(line)
    return {key: "\n".join(value).strip() for key, value in sections.items()}


def parse_use_cases() -> list[dict[str, str]]:
    cases = []
    for path in sorted(USE_CASE_DIR.glob("step-1.3-uc-*.md")):
        text = path.read_text(encoding="utf-8")
        title_match = re.search(r"^#\s+Use Case:\s+(.+)$", text, re.MULTILINE)
        if not title_match:
            raise ValueError(f"Unable to parse use case title from {path}")
        uc_num_match = re.search(r"step-1\.3-uc-(\d+)-", path.name)
        if not uc_num_match:
            raise ValueError(f"Unable to parse use case id from {path}")
        uc_id = f"UC-{int(uc_num_match.group(1)):02d}"
        sections = parse_section_map(text)
        actors = clean_markdown(sections["Actors"])
        primary_match = re.search(r"Primary Actor:\s*(.+)", actors)
        secondary_match = re.search(r"Secondary Actor\(s\):\s*(.+)", actors)
        dependency = clean_markdown(sections["Dependency"])
        if dependency.startswith("- "):
            dependency = dependency[2:]
        cases.append(
            {
                "id": uc_id,
                "name": title_match.group(1).strip(),
                "summary": clean_markdown(sections["Summary"]),
                "dependency": dependency or "None",
                "primary_actor": primary_match.group(1).strip() if primary_match else "",
                "secondary_actors": secondary_match.group(1).strip() if secondary_match else "None",
                "preconditions": clean_markdown(sections["Preconditions"]),
                "main_sequence": clean_markdown(sections["Description of main sequence"]),
                "alternative_sequences": clean_markdown(sections["Description of alternative sequences"]),
                "nonfunctional_requirements": clean_markdown(sections["Nonfunctional Requirements"]),
                "postconditions": clean_markdown(sections["Postcondition"]),
                "outstanding_questions": clean_markdown(sections["Outstanding questions"]),
            }
        )
    return cases


def populate_detail_table(table: Table, use_case: dict[str, str]) -> None:
    row = table.rows[0].cells
    row[0].text = "UC ID and Name:"
    row[1].text = f"{use_case['id']} {use_case['name']}"

    fill_row(table.rows[1], ["Created By:", CREATED_BY, "Date Created:", DATE_CREATED])
    fill_row(
        table.rows[2],
        ["Primary Actor:", use_case["primary_actor"], "Secondary Actors:", use_case["secondary_actors"]],
    )

    merged_rows = [
        ("Summary:", use_case["summary"]),
        ("Preconditions:", use_case["preconditions"]),
        ("Postconditions:", use_case["postconditions"]),
        ("Dependency:", use_case["dependency"]),
        ("Description of main sequence:", use_case["main_sequence"]),
        ("Description of alternative sequences:", use_case["alternative_sequences"]),
        ("Exceptions:", "None explicitly identified in Phase 1."),
        ("Nonfunctional requirements:", use_case["nonfunctional_requirements"]),
        ("Outstanding questions:", use_case["outstanding_questions"] or "None at this stage."),
    ]

    for row_index, (label, value) in enumerate(merged_rows, start=3):
        row = table.rows[row_index].cells
        row[0].text = label
        row[1].text = value


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def format_run_font(run, size: float = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Times New Roman")


def format_paragraph(paragraph: Paragraph) -> None:
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style else ""
    pf = paragraph.paragraph_format

    for run in paragraph.runs:
        format_run_font(run, size=12)

    if style_name == "Heading 1":
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            format_run_font(run, size=14, bold=True)
    elif style_name == "Heading 2":
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            format_run_font(run, size=13, bold=True)
    elif style_name == "Heading 3":
        pf.space_before = Pt(8)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            format_run_font(run, size=12, bold=True)
    elif style_name == "List Paragraph":
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if text and not text.startswith("[Image Placeholder") else WD_ALIGN_PARAGRAPH.LEFT
    else:
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.15
        if paragraph.alignment is None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if text in {"Course Project Report", "Subject: SWD392", "Academic Year: 2026"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(4)
        for run in paragraph.runs:
            format_run_font(run, size=14 if text == "Course Project Report" else 12, bold=(text == "Course Project Report"))

    if text.startswith("[Image Placeholder:"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        for run in paragraph.runs:
            format_run_font(run, size=11, italic=True)
            run.font.color.rgb = RGBColor(90, 90, 90)


def style_table(table: Table, col_widths: list[float] | None = None) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths[: len(row.cells)]):
                set_cell_width(row.cells[idx], width)

    if table.rows:
        header = table.rows[0]
        set_repeat_table_header(header)
        for cell in header.cells:
            set_cell_shading(cell, "D9E2F3")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    format_run_font(run, size=10.5, bold=True)

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index > 0 and row_index % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if row_index > 0 and paragraph.alignment is None:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    format_run_font(run, size=10.5)


def apply_document_formatting(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for paragraph in doc.paragraphs:
        format_paragraph(paragraph)

    table_width_map = {
        0: [1.0, 0.8, 1.2, 3.8],
        1: [2.0, 1.4, 1.7, 1.7],
        2: [0.5, 1.2, 1.8, 3.8],
        3: [0.5, 1.5, 4.8],
        4: [0.8, 2.0, 1.6, 3.0],
        24: [0.6, 1.8, 5.0],
    }

    for idx, table in enumerate(doc.tables):
        style_table(table, table_width_map.get(idx))

        if 5 <= idx <= 22:
            detail_widths = [1.7, 2.45, 1.35, 2.0]
            style_table(table, detail_widths)
            for row_idx in [3, 4, 5, 6, 7, 8, 9, 10, 11]:
                if row_idx < len(table.rows):
                    table.rows[row_idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT



def main() -> None:
    use_cases = parse_use_cases()

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)

    blocks = list(iter_blocks(doc))

    record_table = blocks[28]
    project_name_p = blocks[34]
    project_code_p = blocks[35]
    group_name_p = blocks[36]
    software_type_p = blocks[37]

    problem_paragraphs = [blocks[42], blocks[43], blocks[44]]
    features_intro = blocks[46]
    feature_paragraphs = [blocks[47], blocks[48], blocks[49], blocks[50], blocks[51], blocks[52]]

    context_paragraphs = [blocks[54], blocks[55], blocks[56], blocks[57]]
    nfr_intro = blocks[59]
    nfr_table = blocks[60]

    actor_intro = blocks[64]
    actor_paragraph_2 = blocks[65]
    actor_paragraph_3 = blocks[66]
    actor_paragraph_4 = blocks[67]
    actor_paragraph_5 = blocks[68]
    actor_table = blocks[69]

    use_case_intro = blocks[72]
    use_case_diagram_text = blocks[74]
    use_case_diagram_placeholder = blocks[75]
    use_case_desc_intro = blocks[77]
    use_case_summary_table = blocks[78]
    first_uc_title = blocks[80]
    template_label = blocks[81]
    detail_template_table = blocks[82]
    after_detail_blank_1 = blocks[83]
    after_detail_blank_2 = blocks[84]
    after_detail_blank_3 = blocks[85]

    activity_intro = blocks[87]
    activity_header_text = blocks[88]
    activity_p1 = blocks[89]
    activity_p2 = blocks[90]
    activity_p3 = blocks[91]
    activity_p4 = blocks[92]
    activity_placeholder = blocks[93]

    erd_heading = blocks[94]
    erd_overview = blocks[95]
    erd_subheading = blocks[96]
    erd_intro = blocks[97]

    academic_year = doc.paragraphs[24]
    set_paragraph(academic_year, "Academic Year: 2026", "Normal")

    set_paragraph(project_name_p, f"Project name: {PROJECT_NAME}", "List Paragraph")
    set_paragraph(project_code_p, f"Project code: {PROJECT_CODE}", "List Paragraph")
    set_paragraph(group_name_p, f"Group name: {GROUP_NAME}", "List Paragraph")
    set_paragraph(software_type_p, f"Software type: {SOFTWARE_TYPE}", "List Paragraph")

    fill_row(record_table.rows[1], ["2026-03-10", "A", "Project Team", "Filled Sections I and II with HMSS requirement-modeling content."])

    for paragraph, text in zip(problem_paragraphs, PROBLEM_DESCRIPTION):
        set_paragraph(paragraph, text, "Normal")

    set_paragraph(
        features_intro,
        "The major HMSS features below summarize the business capabilities that drive the functional requirements.",
        "Normal",
    )
    for paragraph, text in zip(feature_paragraphs, MAJOR_FEATURES[:6]):
        set_paragraph(paragraph, text, "Normal")
    last_feature_block = feature_paragraphs[-1]
    for feature in MAJOR_FEATURES[6:]:
        last_feature_block = insert_paragraph_after(last_feature_block, feature, "Normal")

    set_paragraph(
        context_paragraphs[0],
        "At the highest level, HMSS is treated as a single black-box system that exchanges information with human actors and external services.",
        "Normal",
    )
    set_paragraph(
        context_paragraphs[1],
        "External entities: Visitor, Tenant, Owner, System Admin, Google Maps, Cloud Storage, and Email Provider.",
        "Normal",
    )
    set_paragraph(
        context_paragraphs[2],
        "Inbound flows include search criteria, account data, property and listing data, owner verification submissions, rental requests, and administrative control actions.",
        "Normal",
    )
    set_paragraph(
        context_paragraphs[3],
        "[Image Placeholder: Context Diagram]",
        "Normal",
    )
    extra_context = insert_paragraph_after(
        context_paragraphs[3],
        "Outbound flows include search results, room detail information, request-status updates, verification results, listing-control results, notifications, and location display data.",
        "Normal",
    )
    insert_paragraph_after(extra_context, "No include or extend relationships were identified at the requirement-modeling stage.", "Normal")

    set_paragraph(
        nfr_intro,
        "The following nonfunctional requirements summarize the performance, security, availability, reliability, and scalability constraints for HMSS.",
        "Normal",
    )
    ensure_rows(nfr_table, len(NFRS) + 1)
    for row, values in zip(nfr_table.rows[1:], NFRS):
        fill_row(row, values)

    set_paragraph(
        actor_intro,
        "HMSS includes human actors, external-service actors, and one generalized actor used for shared authentication behavior.",
        "Normal",
    )
    set_paragraph(
        actor_paragraph_2,
        "Registered User generalizes Tenant, Owner, and System Admin for the shared Sign In use case.",
        "Normal",
    )
    set_paragraph(actor_paragraph_3, "", "Normal")
    set_paragraph(actor_paragraph_4, "", "Normal")
    set_paragraph(actor_paragraph_5, "", "Normal")

    ensure_rows(actor_table, len(ACTORS) + 1)
    for row, values in zip(actor_table.rows[1:], ACTORS):
        fill_row(row, values)

    set_paragraph(
        use_case_intro,
        "The HMSS use-case model contains 18 requirement-level use cases derived from the approved Phase 1 source material.",
        "Normal",
    )
    set_paragraph(
        use_case_diagram_text,
        "The use case diagram should show Visitor, Registered User, Tenant, Owner, System Admin, Google Maps, Cloud Storage, and Email Provider connected to UC-01 through UC-18.",
        "Normal",
    )
    set_paragraph(use_case_diagram_placeholder, "[Image Placeholder: Use Case Diagram]", "Normal")
    set_paragraph(
        use_case_desc_intro,
        "The summary table and detailed specifications below are populated from the approved HMSS Phase 1 use-case set.",
        "Normal",
    )

    ensure_rows(use_case_summary_table, len(use_cases) + 1)
    for row, use_case in zip(use_case_summary_table.rows[1:], use_cases):
        fill_row(
            row,
            [
                use_case["id"],
                use_case["name"],
                use_case["primary_actor"],
                use_case["summary"],
            ],
        )

    set_paragraph(first_uc_title, f"1. {use_cases[0]['id']} {use_cases[0]['name']}", "Normal")
    set_paragraph(template_label, "", "Normal")
    set_paragraph(after_detail_blank_1, "", "Normal")
    set_paragraph(after_detail_blank_2, "", "Normal")
    set_paragraph(after_detail_blank_3, "", "Normal")
    populate_detail_table(detail_template_table, use_cases[0])

    anchor = detail_template_table
    for idx, use_case in enumerate(use_cases[1:], start=2):
        title_paragraph = insert_paragraph_after(anchor, f"{idx}. {use_case['id']} {use_case['name']}", "Normal")
        detail_table = clone_table_after(detail_template_table, title_paragraph)
        populate_detail_table(detail_table, use_case)
        anchor = detail_table

    set_paragraph(
        activity_intro,
        "The requirement model identifies the following activity-diagram candidates for later visual elaboration.",
        "Normal",
    )
    set_paragraph(activity_header_text, "Recommended activity-diagram coverage:", "Normal")
    set_paragraph(activity_p1, ACTIVITY_ITEMS[0], "List Paragraph")
    set_paragraph(activity_p2, ACTIVITY_ITEMS[1], "List Paragraph")
    set_paragraph(activity_p3, ACTIVITY_ITEMS[2], "List Paragraph")
    set_paragraph(activity_p4, ACTIVITY_ITEMS[3], "List Paragraph")
    set_paragraph(activity_placeholder, "[Image Placeholder: Activity Diagram(s)]", "Normal")

    set_paragraph(erd_heading, "II.6 Entity Relationship Diagram", "Heading 2")
    set_paragraph(
        erd_overview,
        "The core entity model follows the clarified business structure Owner -> Property -> Room Listing, extended with rental-request and verification support entities.",
        "Normal",
    )
    set_paragraph(erd_subheading, "II.6.1 Entity Descriptions", "Heading 3")
    set_paragraph(erd_intro, "The principal entities for the requirement-modeling stage are listed below.", "Normal")

    erd_table = clone_table_after(actor_table, erd_intro)
    ensure_rows(erd_table, len(ERD_ENTITIES) + 1)
    fill_row(erd_table.rows[0], ["#", "Entity", "Description"])
    for row, values in zip(erd_table.rows[1:], ERD_ENTITIES):
        fill_row(row, values)

    rel_heading = insert_paragraph_after(erd_table, "Recommended core relationships:", "Normal")
    rel_anchor = rel_heading
    for relationship in ERD_RELATIONSHIPS:
        rel_anchor = insert_paragraph_after(rel_anchor, relationship, "List Paragraph")
    insert_paragraph_after(rel_anchor, "[Image Placeholder: Entity Relationship Diagram]", "Normal")

    doc.save(OUTPUT_DOCX)
    print(f"Saved filled document to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
